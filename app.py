"""
Consultor Legislativo España – backend FastAPI
Lee los documentos del repositorio legalize-es y expone endpoints de
búsqueda y consulta con IA (Claude).
"""

from __future__ import annotations

import asyncio
import json
import math
import os
import pickle
import re
import threading
import time
import unicodedata
import uuid
from array import array
from collections import Counter
from functools import lru_cache
from pathlib import Path
from typing import List, Optional

import yaml
from fastapi import FastAPI, File, Query, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

import documents
import llm

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

# --- Dónde está el código y dónde están los datos ---------------------------
# Esta aplicación ya no vive dentro del repositorio de datos. legalize-es se
# clona aparte (./setup-corpus.sh) o se indica con CORPUS_PATH, así que hay
# tres rutas distintas y conviene no confundirlas:
#   BASE_DIR    – el código y los estáticos.
#   CORPUS_PATH – el clon de legalize-es, con es/, es-pv/, es-ct/…
#   CACHE_DIR   – los índices derivados. Van junto al código, no dentro del
#                 corpus: ensuciar el clon obliga a ignorarlos en un repo
#                 ajeno y los pierde cualquier `git clean`.
BASE_DIR    = Path(__file__).parent
CORPUS_PATH = Path(os.environ.get("CORPUS_PATH") or BASE_DIR / "legalize-es").expanduser()
CACHE_DIR   = Path(os.environ.get("CACHE_DIR") or BASE_DIR / ".cache").expanduser()
CACHE_DIR.mkdir(parents=True, exist_ok=True)

SEMANTIC_POOL = 60   # candidatos keyword que se re-ordenan semánticamente
INDEX_CACHE = CACHE_DIR / "legalize_index.json"
MAX_RESULTS = 30
PREVIEW_CHARS = 400

# --- Contexto enviado al modelo ---------------------------------------------
# El índice en memoria solo guarda PREVIEW_CHARS del cuerpo de cada norma, que
# son sus primeras líneas: el preámbulo. Con eso es imposible citar un artículo
# concreto — en el Estatuto de los Trabajadores el articulado empieza más allá
# del carácter 130.000. Por eso, al construir el contexto para el modelo se
# releen los ficheros de las normas seleccionadas y se extraen los pasajes que
# coinciden con la consulta, junto al encabezado del artículo al que pertenecen.
CONTEXT_DOCS      = 6     # normas cuyo texto se envía al modelo
FRAG_PER_DOC      = 3     # pasajes por norma
FRAG_WIDTH        = 600   # caracteres por pasaje
# En modo precisión la ventana de Anthropic (1M) permite mucho más contexto.
CONTEXT_DOCS_PREC = 8
FRAG_PER_DOC_PREC = 6
FRAG_WIDTH_PREC   = 1200

# Re-puntuación por texto completo: candidatos de la pasada por palabra clave
# cuyo articulado se relee para reordenarlos. RRF_K amortigua la fusión de
# rangos; cuanto más bajo, más peso tienen las primeras posiciones.
DEEP_POOL = 80
RRF_K     = 10

# --- Índice invertido de texto completo -------------------------------------
# La búsqueda por título no encuentra los grandes códigos en consultas
# temáticas: "estafa" no aparece en el título del Código Penal, ni "permisos"
# en el del Estatuto de los Trabajadores. Este índice recorre el articulado
# una vez al arrancar y guarda qué término aparece en qué norma.
# El coste de E/S ya se pagaba: _parse_doc leía el fichero entero y se
# quedaba solo con los primeros 400 caracteres.
POSTINGS_CACHE  = CACHE_DIR / "legalize_postings.pkl"
# Poda mínima. Al 25% se descartaban términos imprescindibles —"seguridad",
# "personales", "tribunales", "suspensión"— y las consultas se quedaban sin
# nada con lo que emparejar. Ahorraba además poquísimo: los términos muy
# frecuentes son unos cientos, y BM25 ya los degrada solo vía IDF.
FULLTEXT_DF_MAX = 0.90
FULLTEXT_POOL   = 40     # candidatos que aporta este índice a cada búsqueda
# Parámetros BM25. Con solo presencia, una Orden que menciona "vivienda" una
# vez puntuaba igual que la Ley de Arrendamientos Urbanos, que la menciona
# cientos: sin frecuencia de término no hay forma de distinguirlas. K1 gradúa
# cuánto aporta repetir un término y B cuánto se penaliza la longitud (si no,
# los códigos ganarían siempre por ser largos).
BM25_K1 = 1.5
BM25_B  = 0.75
# Mínimo para que el texto cuente: 3 términos distintos en un mismo pasaje, o
# una locución contigua (que vale 2) más un término suelto.
BODY_MIN_SCORE = 3.0
BODY_WEIGHT    = 0.5

# ---------------------------------------------------------------------------
# Aprendizaje por Refuerzo / Retroalimentación Adaptativa de Búsqueda
# ---------------------------------------------------------------------------
LEARNING_RATE: float = 0.25  # Hiperparámetro alfa de aprendizaje por refuerzo
_reinforcement_scores: dict[str, float] = {}
_reinforcement_lock = threading.Lock()


def _record_reinforcement_feedback(doc_id: str, reward: float = 1.0) -> None:
    """Acumula recompensa por aprendizaje por refuerzo para normas citadas con éxito."""
    if not doc_id:
        return
    with _reinforcement_lock:
        current = _reinforcement_scores.get(doc_id, 0.0)
        _reinforcement_scores[doc_id] = min(50.0, current + reward)

REGIONS: dict[str, str] = {
    "es": "España (Estatal)",
    "es-an": "Andalucía",
    "es-ar": "Aragón",
    "es-as": "Asturias",
    "es-cb": "Cantabria",
    "es-cl": "Castilla y León",
    "es-cm": "Castilla-La Mancha",
    "es-cn": "Canarias",
    "es-ct": "Cataluña",
    "es-ex": "Extremadura",
    "es-ga": "Galicia",
    "es-ib": "Islas Baleares",
    "es-mc": "Murcia",
    "es-md": "Madrid",
    "es-nc": "Navarra",
    "es-pv": "País Vasco",
    "es-ri": "La Rioja",
    "es-vc": "Valencia",
}

RANK_LABELS: dict[str, str] = {
    "ley": "Ley",
    "ley_organica": "Ley Orgánica",
    "real_decreto_ley": "Real Decreto-ley",
    "real_decreto": "Real Decreto",
    "decreto": "Decreto",
    "decreto_ley": "Decreto-ley",
    "decreto_legislativo": "Decreto Legislativo",
    "orden": "Orden",
    "orden_ministerial": "Orden Ministerial",
    "resolucion": "Resolución",
    "instruccion": "Instrucción",
    "circular": "Circular",
    "convenio": "Convenio",
    "acuerdo": "Acuerdo",
    "anuncio": "Anuncio",
    "correccion_errores": "Corrección de errores",
    "reglamento": "Reglamento",
    "estatuto": "Estatuto",
    "carta": "Carta",
    "tratado": "Tratado",
}

STOP_WORDS = {
    "de", "la", "el", "en", "y", "a", "que", "los", "las", "se", "del",
    "al", "por", "con", "un", "una", "para", "es", "su", "sus", "o", "e",
    "no", "si", "más", "mas", "sobre", "esta", "este", "ello", "como", "entre",
    "lo", "le", "les", "todo", "todos", "todas", "ha", "han", "hay", "ser",
    "fue", "son", "cual", "cuales",
}

# Legal abbreviations → expanded query terms for better recall
ABBREVIATIONS: dict[str, str] = {
    # Impuestos
    "irpf": "impuesto renta personas fisicas",
    "iva": "impuesto valor añadido",
    "is": "impuesto sociedades",
    "ibi": "impuesto bienes inmuebles",
    "iae": "impuesto actividades economicas",
    "isd": "impuesto sucesiones donaciones",
    "ip": "impuesto patrimonio",
    "irnr": "impuesto renta no residentes",
    "itpajd": "impuesto transmisiones patrimoniales actos juridicos documentados",
    # Leyes y códigos
    "lau": "ley arrendamientos urbanos alquiler vivienda",
    "lsc": "ley sociedades capital mercantil",
    "trlsc": "texto refundido ley sociedades capital",
    "et": "estatuto trabajadores laboral empleo",
    "lgss": "ley general seguridad social pensiones",
    "lgt": "ley general tributaria hacienda",
    "ce": "constitucion española derechos fundamentales",
    "cc": "codigo civil",
    "cp": "codigo penal delito",
    "lec": "ley enjuiciamiento civil proceso judicial",
    "lecrim": "ley enjuiciamiento criminal penal",
    "lgp": "ley general presupuestos generales estado",
    "lpac": "ley procedimiento administrativo comun administracion publica",
    "lrjsp": "regimen juridico sector publico administracion",
    "lopd": "ley proteccion datos personales privacidad",
    "rgpd": "reglamento general proteccion datos personales privacidad",
    "lopdgdd": "proteccion datos garantia derechos digitales",
    "loe": "ley organica educacion escolar",
    "lomce": "ley organica mejora calidad educativa",
    "lomloe": "ley organica modificacion educacion",
    # Prestaciones / situaciones
    "erte": "expediente regulacion temporal empleo",
    "ere": "expediente regulacion empleo",
    "imv": "ingreso minimo vital renta",
    "smi": "salario minimo interprofesional",
    "pie": "participacion ingresos estado financiacion",
    "dana": "depresion aislada niveles altos catastrofe inundacion emergencia",
    # Organismos
    "aeat": "agencia tributaria hacienda",
    "boe": "boletin oficial estado",
    "tc": "tribunal constitucional",
    "ts": "tribunal supremo",
    "tjue": "tribunal justicia union europea",
    "cedh": "convenio europeo derechos humanos",
    "ss": "seguridad social",
    "sepe": "servicio empleo publico estatal",
}


# ---------------------------------------------------------------------------
# Documentos adjuntos
# ---------------------------------------------------------------------------
# El usuario puede adjuntar una resolución, un contrato o un informe para que
# la consulta se base en él. El texto extraído se guarda SOLO EN MEMORIA y
# caduca: son documentos con datos personales (nombre, DNI, cuentas) y escribir
# su contenido en /tmp dejaría un rastro que nadie ha pedido.
ADJUNTO_MAX_BYTES = 10 * 1024 * 1024   # 10 MB por fichero
ADJUNTO_TTL       = 3600               # 1 hora
ADJUNTO_MAX_CHARS = 400_000            # tope de texto extraído
ADJUNTO_FRAGS     = 6                  # pasajes del adjunto que ve el modelo
ADJUNTO_WIDTH     = 700

_adjuntos: dict = {}
_adjuntos_lock = threading.Lock()


# ---------------------------------------------------------------------------
# Locuciones coloquiales → redacción legal
# ---------------------------------------------------------------------------
# ABBREVIATIONS resuelve acrónimos palabra a palabra, pero no locuciones. Y el
# problema real de las consultas ciudadanas es de vocabulario: se pregunta por
# "brecha de datos" cuando el BOE dice "violación de la seguridad de los datos
# personales". Sin esta traducción el buscador solo puede acertar de rebote —
# "brecha" aparece en una única norma de todo el corpus estatal.
#
# Los valores van CON tildes a propósito: el emparejamiento se hace por
# subcadena sobre el texto real de las normas, y "violacion" no casa con
# "violación".
LEGAL_TERMS: dict[str, str] = {
    # Protección de datos
    "brecha de datos":            "violación de la seguridad de los datos personales",
    "brecha de seguridad":        "violación de la seguridad de los datos personales",
    "fuga de datos":              "violación de la seguridad de los datos personales",
    "filtracion de datos":        "violación de la seguridad de los datos personales",
    "derecho al olvido":          "supresión de datos personales",
    "camaras de seguridad":       "tratamiento de datos con fines de vigilancia",
    "listas de morosos":          "sistemas de información crediticia solvencia patrimonial",
    # Laboral y Seguridad Social
    "pension de jubilacion":      "pensión de jubilación carencia periodo minimo cotizacion Ley General Seguridad Social TRLGSS",
    "jubilacion discapacidad":    "pensión de jubilación personas con discapacidad carencia no contributiva Ley General Seguridad Social TRLGSS",
    "pension no contributiva":    "pensión no contributiva invalidez jubilación imserso carencia Ley General Seguridad Social TRLGSS",
    "menos de 15 años cotizados": "periodo minimo de cotizacion carencia quince años pensión no contributiva TRLGSS",
    "baja laboral":               "incapacidad temporal",
    "baja por maternidad":        "suspensión del contrato por nacimiento y cuidado de menor",
    "baja por paternidad":        "suspensión del contrato por nacimiento y cuidado de menor",
    "el paro":                    "prestación por desempleo",
    "cobrar el paro":             "prestación por desempleo",
    "teletrabajo":                "trabajo a distancia",
    "acoso laboral":              "acoso en el trabajo",
    "horas extra":                "horas extraordinarias",
    # Vivienda y consumo
    "alquiler":                   "arrendamiento",
    "clausula suelo":             "cláusulas abusivas préstamo hipotecario",
    "okupas":                     "usurpación de bienes inmuebles",
    # Familia
    "custodia de los hijos":      "guarda y custodia",
    "pension de alimentos":       "alimentos entre parientes",
    # Tráfico
    "multa de trafico":           "infracciones en materia de tráfico y seguridad vial",
    "carnet por puntos":          "permiso de conducción por puntos",
}

# Resueltas de más larga a más corta para que la locución más específica gane.
_LEGAL_TERMS_SORTED = sorted(LEGAL_TERMS.items(), key=lambda kv: -len(kv[0]))


def _normalize_text(text: str) -> str:
    """Normaliza texto a minúsculas y elimina tildes/diacríticos para insensibilidad de acentos."""
    if not text:
        return ""
    nfkd = unicodedata.normalize("NFD", text.lower())
    return "".join(c for c in nfkd if unicodedata.category(c) != "Mn")


_LEGAL_TERMS_SORTED_NORM = [(_normalize_text(k), _normalize_text(v)) for k, v in _LEGAL_TERMS_SORTED]

# ---------------------------------------------------------------------------
# Estado global del índice
# ---------------------------------------------------------------------------

_index: List[dict] = []
_index_ready = threading.Event()
_index_error: str | None = None
_index_total = 0
_index_loaded = 0
# Índice invertido del articulado. Cada entrada es un array('i') con pares
# (id_de_norma, veces_que_aparece) intercalados: 8 bytes por posting, frente a
# los ~60 que costaría una lista de tuplas de Python.
_postings: dict = {}
_doc_len: array = array("i")
_avg_len = 1.0
_n_docs = 0


# ---------------------------------------------------------------------------
# Parseo de documentos
# ---------------------------------------------------------------------------

def _parse_doc(filepath: Path) -> tuple | None:
    """Devuelve (metadatos, cuerpo_limpio) o None si el fichero no es válido.

    El cuerpo se devuelve aparte y NO se guarda en los metadatos: estos se
    serializan a la caché JSON del índice, y meter ahí 964 MB de articulado
    la haría inservible. Solo se usa para alimentar el índice invertido.
    """
    try:
        text = filepath.read_text(errors="ignore")
    except OSError:
        return None

    if not text.startswith("---"):
        return None

    end = text.find("---", 3)
    if end == -1:
        return None

    yaml_text = text[3:end].strip()
    content = text[end + 3:].strip()

    try:
        meta = yaml.safe_load(yaml_text)
    except yaml.YAMLError:
        return None

    if not isinstance(meta, dict):
        return None

    # Strip markdown from content preview
    clean = re.sub(r"#+\s*", "", content)
    clean = re.sub(r"\*+", "", clean)
    clean = re.sub(r"\n{2,}", "\n", clean).strip()

    meta["_preview"] = clean[:PREVIEW_CHARS]
    meta["_region"] = filepath.parent.name
    meta["_filepath"] = str(filepath)
    meta["_filename"] = filepath.name

    # Normalise subjects to list
    subjects = meta.get("subjects")
    if isinstance(subjects, str):
        meta["subjects"] = [s.strip() for s in subjects.split(",")]
    elif not isinstance(subjects, list):
        meta["subjects"] = []

    # Ensure dates are converted to strings so JSON serialization & comparison never fail
    for date_key in ("publication_date", "last_updated"):
        if date_key in meta and meta[date_key] is not None:
            meta[date_key] = str(meta[date_key])

    return meta, clean


def _install_postings(postings: dict, doc_len: array, n_docs: int) -> None:
    """Publica el índice invertido descartando los términos que no discriminan.

    Palabras como "artículo", "ley" o "disposición" salen en casi todas las
    normas: ocupan memoria, alargan cada consulta y no aportan señal.
    """
    global _postings, _doc_len, _avg_len, _n_docs
    if not n_docs:
        return
    tope = max(1, int(n_docs * FULLTEXT_DF_MAX))
    # len(p)//2 porque cada posting ocupa dos huecos: id y frecuencia.
    _postings = {t: p for t, p in postings.items() if len(p) // 2 <= tope}
    _doc_len = doc_len
    total = sum(doc_len)
    _avg_len = (total / n_docs) if n_docs and total else 1.0
    _n_docs = n_docs


def _fulltext_scores(terms: List[str], limit: int) -> dict:
    """Normas cuyo articulado menciona los términos, puntuadas con BM25.

    Frente a contar solo presencia, BM25 tiene en cuenta cuántas veces aparece
    cada término y la longitud de la norma. Es lo que separa a la Ley de
    Arrendamientos Urbanos, que habla de "arrendamiento" en cada artículo, de
    las 1.395 normas que lo mencionan de pasada.
    """
    if not _postings or not _n_docs:
        return {}

    acumulado: dict = {}
    for term in terms:
        lista = _postings.get(term)
        if not lista:
            continue
        df = len(lista) // 2
        idf = math.log(1 + (_n_docs - df + 0.5) / (df + 0.5))
        for i in range(0, len(lista), 2):
            doc_id = lista[i]
            tf = lista[i + 1]
            largo = _doc_len[doc_id] if doc_id < len(_doc_len) else 0
            norm = BM25_K1 * (1 - BM25_B + BM25_B * (largo / _avg_len))
            acumulado[doc_id] = acumulado.get(doc_id, 0.0) + \
                idf * (tf * (BM25_K1 + 1)) / (tf + norm)

    if len(acumulado) <= limit:
        return acumulado
    return dict(sorted(acumulado.items(), key=lambda kv: -kv[1])[:limit])


def _build_index() -> None:
    global _index, _index_error, _index_total, _index_loaded

    if not CORPUS_PATH.is_dir():
        _index_error = (
            f"No se encuentra el corpus en {CORPUS_PATH}. Ejecuta ./setup-corpus.sh "
            f"para clonar legalize-es, o define CORPUS_PATH apuntando a tu copia."
        )
        _index_ready.set()
        return

    dirs = sorted(
        [d for d in CORPUS_PATH.iterdir() if d.is_dir() and d.name.startswith("es")],
        key=lambda d: (0 if d.name == "es" else 1, d.name),
    )

    if not dirs:
        _index_error = (
            f"{CORPUS_PATH} existe pero no contiene directorios de jurisdicción "
            f"(es/, es-pv/, es-ct/…). ¿Apunta CORPUS_PATH a la raíz de legalize-es?"
        )
        _index_ready.set()
        return

    all_files: List[Path] = []
    for d in dirs:
        all_files.extend(sorted(d.glob("*.md")))

    # Comprobar validez de caché contra la fecha de modificación de los ficheros markdown
    if INDEX_CACHE.exists() and POSTINGS_CACHE.exists():
        try:
            cache_mtime = INDEX_CACHE.stat().st_mtime
            latest_md_mtime = max((fp.stat().st_mtime for fp in all_files), default=0.0)
            if cache_mtime >= latest_md_mtime:
                data = json.loads(INDEX_CACHE.read_text(encoding="utf-8"))
                with POSTINGS_CACHE.open("rb") as fh:
                    n_docs, postings, doc_len = pickle.load(fh)
                # El índice guarda rutas absolutas en "_filepath". Si el corpus
                # se movió o CORPUS_PATH apunta a otra copia, la caché es válida
                # en tamaño pero sus rutas ya no existen, y la comparación de
                # mtime no lo detecta. Se comprueba explícitamente.
                mismo_corpus = str(data[0].get("_filepath", "")).startswith(str(CORPUS_PATH)) if data else False
                if (isinstance(data, list) and len(data) > 100
                        and n_docs == len(data) == len(doc_len)
                        and mismo_corpus):
                    _index = data
                    _install_postings(postings, doc_len, n_docs)
                    _index_ready.set()
                    return
        except Exception:
            pass

    _index_total = len(all_files)

    docs: List[dict] = []
    postings: dict = {}
    doc_len = array("i")

    for i, fp in enumerate(all_files):
        parsed = _parse_doc(fp)
        if parsed:
            doc, body = parsed
            doc_id = len(docs)
            docs.append(doc)
            # Aquí NO se puede usar _tokenise: deduplica, y BM25 necesita
            # justamente cuántas veces aparece cada término.
            tokens = _tokenise_all(body)
            doc_len.append(len(tokens))
            for term, veces in Counter(tokens).items():
                lista = postings.get(term)
                if lista is None:
                    lista = postings[term] = array("i")
                lista.append(doc_id)
                lista.append(veces)
        _index_loaded = i + 1

    _index = docs
    _install_postings(postings, doc_len, len(docs))

    try:
        tmp_index = INDEX_CACHE.with_suffix(".tmp")
        tmp_postings = POSTINGS_CACHE.with_suffix(".tmp")
        tmp_index.write_text(json.dumps(docs, ensure_ascii=False), encoding="utf-8")
        with tmp_postings.open("wb") as fh:
            pickle.dump((len(docs), _postings, _doc_len), fh,
                        protocol=pickle.HIGHEST_PROTOCOL)
        os.replace(tmp_index, INDEX_CACHE)
        os.replace(tmp_postings, POSTINGS_CACHE)
    except Exception:
        pass
    finally:
        _index_ready.set()


threading.Thread(target=_build_index, daemon=True, name="index-builder").start()


# ---------------------------------------------------------------------------
# Motor de búsqueda
# ---------------------------------------------------------------------------

def _tokenise_all(text: str) -> List[str]:
    """Todos los tokens útiles, CON repeticiones y SIN tildes. Para indexar con BM25."""
    clean = _normalize_text(text)
    return [t for t in re.findall(r"[a-z0-9ñ]+", clean)
            if t not in STOP_WORDS and len(t) > 2]


def _tokenise(text: str) -> List[str]:
    clean = _normalize_text(text)
    tokens = re.findall(r"[a-z0-9ñ]+", clean)
    vistos: set = set()
    unicos: List[str] = []
    for t in tokens:
        if t in STOP_WORDS or len(t) < 2 or t in vistos:
            continue
        vistos.add(t)
        unicos.append(t)
    return unicos


def _expand_query(query: str) -> str:
    """Traduce la consulta al vocabulario con el que está redactado el BOE, insensible a acentos."""
    lowered = _normalize_text(query)

    for phrase, canonical in _LEGAL_TERMS_SORTED_NORM:
        if phrase in lowered:
            lowered = lowered.replace(phrase, canonical)

    expanded: List[str] = []
    for word in lowered.split():
        key = re.sub(r"[^a-z0-9ñ]", "", word)
        expanded.append(ABBREVIATIONS.get(key, word))
    return " ".join(expanded)


# Encabezados que dan nombre al pasaje: "###### Artículo 37. Descanso semanal."
_HEADING_RE = re.compile(
    r"^#{1,6}\s*((?:Artículo|Disposición|Sección|Capítulo|Título)\b[^\n]*)",
    re.IGNORECASE,
)


@lru_cache(maxsize=1024)
def _analyse_body(filepath: str, terms_key: tuple, title: str) -> tuple:
    """Puntúa el articulado de una norma frente a los términos de la consulta.

    Devuelve (puntuación_global, pasajes), donde cada pasaje es
    (puntuación, orden, encabezado, texto). Es la base tanto de la re-búsqueda
    por texto completo como de los fragmentos que se envían al modelo, por eso
    va cacheada: en una misma consulta se pide dos veces por documento.
    """
    try:
        text = Path(filepath).read_text(errors="ignore")
    except OSError:
        return (0.0, ())
    return _analyse_text(text, terms_key, title)


def _analyse_text(text: str, terms_key: tuple, title: str = "") -> tuple:
    """Mismo análisis, pero sobre texto en memoria.

    Se separa de _analyse_body para poder aplicarlo también a los documentos
    que adjunta el usuario, que no están en el corpus ni tienen fichero.
    """
    terms = list(terms_key)

    # Los términos que ya aparecen en el título de la norma sirvieron para
    # elegirla, pero dentro de ella no discriminan nada: "estatuto" o
    # "trabajadores" salen en su artículo de aprobación y arrastrarían
    # pasajes puramente formales. Aquí solo interesan los que localizan.
    # Comparando por palabras completas, no por subcadenas: el título del
    # Estatuto dice "por el que se aprueba", y "aprueba" contiene "prueba",
    # que descartaría justo el término buscado en "periodo de prueba".
    if title:
        title_terms = set(_tokenise(title))
        terms = [t for t in terms if t not in title_terms]
    if not terms:
        return (0.0, ())

    # Saltar el frontmatter YAML: aquí solo interesa el articulado.
    body = text
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            body = text[end + 3:]

    # Primera pasada: separar encabezados de párrafos de texto.
    heading = ""
    content: List[tuple] = []
    for order, raw in enumerate(body.split("\n")):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            # Solo los encabezados de articulado sirven para citar; los demás
            # (título de la norma, capítulos sin número) no deben puntuar.
            match = _HEADING_RE.match(line)
            if match:
                heading = match.group(1).strip()
            continue
        content.append((order, heading, line, _normalize_text(line)))

    if not content:
        return (0.0, ())

    # Frecuencia de cada término dentro de esta norma. Sin esto, palabras
    # ubicuas como "días" o "trabajadores" pesan lo mismo que el término que
    # de verdad discrimina ("fallecimiento") y sepultan el párrafo correcto.
    freq = {t: sum(1 for _, _, _, low in content if t in low) for t in terms}

    # Muchas consultas jurídicas son locuciones ("periodo de prueba",
    # "despido improcedente"): sus palabras sueltas son comunes, pero juntas
    # identifican el artículo exacto. Se premia que aparezcan contiguas,
    # admitiendo una palabra intermedia ("periodo DE prueba").
    pares = []
    for a, b in zip(terms, terms[1:]):
        if freq.get(a) and freq.get(b):
            pares.append((
                re.compile(rf"{re.escape(a)}\w*\W+(?:\w+\W+)?{re.escape(b)}"),
                1.0 / freq[a] + 1.0 / freq[b],
            ))

    def relevancia(texto: str) -> float:
        """Cuanto más raro es el término dentro de esta norma, más pesa."""
        score = sum(1.0 / freq[t] for t in terms if freq.get(t) and t in texto)
        score += sum(peso * 3 for patron, peso in pares if patron.search(texto))
        return score

    # El epígrafe del artículo dice de qué trata ("Artículo 37. Descanso
    # semanal, fiestas y permisos"). Sin él, un párrafo suelto sobre
    # fallecimiento en el artículo de suspensión del contrato compite de igual
    # a igual con el que regula los días de permiso, que es el que se busca.
    head_scores: dict = {}
    scored: List[tuple] = []
    for order, head, line, low in content:
        if head not in head_scores:
            head_scores[head] = relevancia(_normalize_text(head)) if head else 0.0
        score = relevancia(low) + 0.5 * head_scores[head]
        if score:
            scored.append((score, order, head, line))

    if not scored:
        return (0.0, ())

    ordenados = sorted(scored, key=lambda s: (-s[0], s[1]))

    # OJO: la puntuación de arriba pondera por rareza DENTRO de esta norma, así
    # que no es comparable entre normas distintas — un término que sale una vez
    # en una Orden de veinte líneas vale 1,0 y en el Estatuto de los
    # Trabajadores 0,006. Sirve para elegir el pasaje, no para comparar
    # documentos. Para eso se usa una medida sin escala: cuántos términos
    # distintos de la consulta reúne el mejor pasaje, más las locuciones que
    # aparecen contiguas en él.
    total = 0.0
    for _, _, head, line in ordenados[:5]:
        texto = _normalize_text(f"{head} {line}")
        cobertura = sum(1 for t in terms if t in texto)
        locuciones = sum(1 for patron, _ in pares if patron.search(texto))
        total = max(total, cobertura + 2.0 * locuciones)
    # Solo se conservan los mejores: el resto nunca se usa y el Estatuto de
    # los Trabajadores solo aportaría miles de líneas a la caché.
    return (total, tuple(ordenados[:30]))


def _relevant_fragments(filepath: str, terms: List[str],
                        max_frags: int, width: int,
                        title: str = "") -> List[str]:
    """Pasajes de una norma que tocan la consulta, listos para el modelo.

    Cada pasaje va precedido del encabezado del artículo en el que aparece,
    que es lo que permite al modelo citar "artículo 37.3.b" en vez de
    aproximar un número.
    """
    _, scored = _analyse_body(filepath, tuple(terms), title)
    if not scored:
        return []

    # Mejores por relevancia; después se reordenan para que el modelo los lea
    # en el orden en que aparecen en la norma.
    best = sorted(scored[:max_frags], key=lambda s: s[1])

    fragments: List[str] = []
    for _, _, head, line in best:
        cut = line[:width]
        if len(line) > width:
            cut = cut.rsplit(" ", 1)[0] + "…"
        fragments.append(f"{head}\n{cut}" if head else cut)
    return fragments


def _preprocesar_imagen_pil(img):
    """Normaliza y escala imágenes para maximizar la tasa de éxito de Tesseract OCR."""
    from PIL import Image
    # 1. Aplanar transparencias (RGBA/LA/P) sobre fondo blanco sólido
    if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
        bg = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode == "P":
            img = img.convert("RGBA")
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")

    # 2. Reescalado automático si la resolución es baja (< 1000px de ancho)
    if img.width < 1000:
        factor = 1000.0 / float(img.width)
        new_size = (int(img.width * factor), int(img.height * factor))
        img = img.resize(new_size, Image.Resampling.LANCZOS)
    return img


def _ocr_imagen_pil(img) -> str:
    """Ejecuta Tesseract OCR sobre una imagen PIL con estrategias multimodales de PSM."""
    import pytesseract
    img = _preprocesar_imagen_pil(img)

    # 1. Intento primario: PSM por defecto
    try:
        t = pytesseract.image_to_string(img, lang="spa+eng").strip()
        if len(t) > 5:
            return t
    except Exception:
        pass

    # 2. Intento secundario: PSM 6 (Bloque uniforme de texto)
    try:
        t = pytesseract.image_to_string(img, lang="spa+eng", config="--psm 6").strip()
        if len(t) > 5:
            return t
    except Exception:
        pass

    # 3. Intento terciario: PSM 11 (Texto disperso / capturas de pantalla)
    try:
        t = pytesseract.image_to_string(img, lang="spa+eng", config="--psm 11").strip()
        if len(t) > 5:
            return t
    except Exception:
        pass

    # 4. Fallback sólo castellano
    try:
        return pytesseract.image_to_string(img, lang="spa").strip()
    except Exception:
        return ""


def _extraer_texto(nombre: str, datos: bytes) -> str:
    """Saca el texto de CUALQUIER formato de archivo (PDF, DOCX, imágenes OCR, HTML, TXT, o binarios)."""
    ext = Path(nombre or "").suffix.lower()

    # 1. PDFs (Texto vectorial + OCR Tesseract doble vía)
    if ext == ".pdf":
        texto_pdf = ""
        try:
            import io
            from pypdf import PdfReader
            lector = PdfReader(io.BytesIO(datos))
            partes = []
            for pagina in lector.pages:
                try:
                    partes.append(pagina.extract_text() or "")
                except Exception:
                    continue
            texto_pdf = "\n".join(partes).strip()
            if len(texto_pdf) >= 50:
                return texto_pdf
        except Exception:
            texto_pdf = ""

        # OCR Fallback
        partes_ocr = []
        try:
            from pdf2image import convert_from_bytes
            imagenes = convert_from_bytes(datos, dpi=150, first_page=1, last_page=15)
            for img in imagenes:
                t = _ocr_imagen_pil(img)
                if t and len(t.strip()) > 10:
                    partes_ocr.append(t.strip())
        except Exception:
            pass

        if not partes_ocr:
            try:
                import io
                from PIL import Image
                from pypdf import PdfReader
                lector_img = PdfReader(io.BytesIO(datos))
                for pag in lector_img.pages[:15]:
                    for img_obj in pag.images:
                        try:
                            pil_img = Image.open(io.BytesIO(img_obj.data))
                            txt_img = _ocr_imagen_pil(pil_img)
                            if txt_img and len(txt_img.strip()) > 10:
                                partes_ocr.append(txt_img.strip())
                        except Exception:
                            pass
            except Exception:
                pass

        if partes_ocr:
            return "\n".join(partes_ocr)
        if texto_pdf:
            return texto_pdf

    # 2. Documentos Word (DOCX / DOC)
    if ext in (".docx", ".doc"):
        try:
            import io
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(datos))
            partes = [p.text for p in doc.paragraphs]
            for tabla in doc.tables:
                for fila in tabla.rows:
                    partes.append(" | ".join(c.text for c in fila.cells))
            txt_docx = "\n".join(partes).strip()
            if txt_docx:
                return txt_docx
        except Exception:
            pass

    # 3. Imágenes conocidas (PNG, JPG, JPEG, TIFF, BMP, WEBP, PNM, PGM, PPM)
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp", ".pnm", ".pgm", ".ppm"):
        try:
            import io
            from PIL import Image
            img = Image.open(io.BytesIO(datos))
            return _ocr_imagen_pil(img)
        except Exception:
            pass

    # 4. Intento Universal de OCR con Pillow (si el archivo es una imagen con extensión arbitraria)
    try:
        import io
        from PIL import Image
        img = Image.open(io.BytesIO(datos))
        txt_ocr = _ocr_imagen_pil(img)
        if len(txt_ocr) > 5:
            return txt_ocr
    except Exception:
        pass

    # 5. Decodificación de Texto Universal (UTF-8, Latin-1, CP1252, ISO-8859)
    for enc in ("utf-8", "latin-1", "cp1252", "iso-8859-15"):
        try:
            decoded = datos.decode(enc)
            clean_txt = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", decoded).strip()
            if len(clean_txt) > 5:
                return clean_txt
        except UnicodeDecodeError:
            continue

    # 6. Extractor Universal de cadenas legibles (fallback de último recurso para archivos binarios desconocidos)
    matches = re.findall(r"[a-zA-Z0-9áéíóúÁÉÍÓÚñÑ\s.,;:/\-–—_()]{4,}", datos.decode("latin-1", errors="ignore"))
    clean_binary = "\n".join(m.strip() for m in matches if len(m.strip()) > 3).strip()
    return clean_binary


def _purgar_adjuntos(max_items: int = 50) -> None:
    """Elimina los adjuntos caducados y limita el tamaño máximo en memoria."""
    limite = time.time() - ADJUNTO_TTL
    with _adjuntos_lock:
        # Purga por antigüedad (1 hora)
        for clave in [k for k, v in _adjuntos.items() if v["ts"] < limite]:
            _adjuntos.pop(clave, None)
        # Purga por exceso de capacidad (máximo 50 entradas más recientes)
        if len(_adjuntos) > max_items:
            ordenados = sorted(_adjuntos.items(), key=lambda kv: kv[1]["ts"])
            for clave, _ in ordenados[: len(_adjuntos) - max_items]:
                _adjuntos.pop(clave, None)



def _terminos_salientes(texto: str, tope: int = 12) -> List[str]:
    """Términos que mejor caracterizan un adjunto, para enriquecer la búsqueda.

    Se comparan sus frecuencias con las del corpus: una palabra que aquí sale
    mucho y en el BOE poco es la que identifica de qué trata el documento.
    """
    if not texto:
        return []
    tokens = _tokenise_all(texto)
    # Los números puros (años, números de expediente, importes) no localizan
    # normativa: "2026" no lleva a la ley de tráfico.
    cuentas = Counter(t for t in tokens if not t.isdigit())
    if not cuentas or not _postings or not _n_docs:
        return [t for t, _ in cuentas.most_common(tope)]

    # En un documento corto —una resolución de una página— cada término
    # relevante aparece una sola vez. Exigir dos repeticiones dejaba fuera
    # justo las palabras que identifican la materia.
    minimo = 2 if len(tokens) > 600 else 1

    puntuados = []
    for termino, veces in cuentas.items():
        if veces < minimo:
            continue
        lista = _postings.get(termino)
        df = (len(lista) // 2) if lista else 1
        puntuados.append((veces * math.log(1 + _n_docs / max(df, 1)), termino))
    puntuados.sort(reverse=True)
    return [t for _, t in puntuados[:tope]]


def _pasajes_adjunto(texto: str, terms: List[str],
                     max_frags: int = ADJUNTO_FRAGS,
                     width: int = ADJUNTO_WIDTH) -> str:
    """Trozos del adjunto relacionados con la consulta, listos para el prompt.

    Una resolución de veinte páginas no cabe en la ventana de 8.192 tokens del
    modelo local, así que se envían solo los pasajes pertinentes. Si ninguno
    coincide, se manda el principio: en un escrito administrativo el
    encabezamiento ya identifica órgano, expediente y objeto.
    """
    if not texto:
        return ""
    _, pasajes = _analyse_text(texto, tuple(terms))
    elegidos = sorted(pasajes[:max_frags], key=lambda s: s[1])
    if not elegidos:
        return f"    {texto[:width * 2]}"
    return "\n\n".join(
        f"    {(h + chr(10)) if h else ''}{l[:width]}" for _, _, h, l in elegidos)


def _score_doc(doc: dict, terms: List[str]) -> int:
    title = _normalize_text(doc.get("title") or "")
    subjects = _normalize_text(" ".join(doc.get("subjects") or []))
    preview = _normalize_text(doc.get("_preview") or "")
    dept = _normalize_text(doc.get("department") or "")
    rank = _normalize_text(doc.get("rank") or "")
    alerts = _normalize_text(doc.get("alerts") or "")

    score = 0
    for term in terms:
        if term in title:
            score += 12
        if term in subjects:
            score += 8
        if term in dept:
            score += 5
        if term in alerts:
            score += 4
        if term in rank:
            score += 3
        if term in preview:
            score += 2

    # Bonificación estructural a los Textos Refundidos, Códigos estatales y Leyes Orgánicas
    # frente a Reales Decretos-leyes de modificación puntual.
    if score > 0:
        if "texto refundido" in title or "codigo" in title:
            score += 15
        elif "ley organica" in title or doc.get("rank") in ("ley_organica", "real_decreto_legislativo"):
            score += 10

        # Aprendizaje por Refuerzo: Impulso dinámico adaptativo por feedback exitoso acumulado
        doc_id = doc.get("identifier") or doc.get("_filename")
        if doc_id and doc_id in _reinforcement_scores:
            rl_boost = int(_reinforcement_scores[doc_id] * LEARNING_RATE)
            score += rl_boost

    return score


def search(
    query: str,
    region: str | None = None,
    rank_filter: str | None = None,
    status_filter: str | None = None,
    limit: int = MAX_RESULTS,
    deep: bool = True,
) -> List[dict]:
    terms = _tokenise(_expand_query(query))
    if not terms:
        return []

    results: List[tuple[int, dict]] = []

    for doc in _index:
        if region and doc.get("_region") != region:
            continue
        if rank_filter and doc.get("rank") != rank_filter:
            continue
        if status_filter and doc.get("status") != status_filter:
            continue

        score = _score_doc(doc, terms)
        if score > 0:
            results.append((score, doc))

    results.sort(key=lambda x: (x[0], x[1].get("publication_date") or ""), reverse=True)

    if not deep:
        return [doc for _, doc in results[:limit]]

    # --- Aporte del índice invertido ---------------------------------------
    # Los grandes códigos no llevan el asunto en el título, así que la pasada
    # anterior nunca los encuentra. Aquí entran por su articulado.
    def pasa_filtros(doc: dict) -> bool:
        return not (
            (region and doc.get("_region") != region)
            or (rank_filter and doc.get("rank") != rank_filter)
            or (status_filter and doc.get("status") != status_filter)
        )

    ft = _fulltext_scores(terms, FULLTEXT_POOL)
    rank_texto_completo: dict = {}
    ya = {id(doc) for _, doc in results[:DEEP_POOL]}
    extra: List[tuple] = []
    for pos, (doc_id, _) in enumerate(sorted(ft.items(), key=lambda kv: -kv[1])):
        if doc_id >= len(_index):
            continue
        doc = _index[doc_id]
        if not pasa_filtros(doc):
            continue
        rank_texto_completo[id(doc)] = pos
        if id(doc) not in ya:
            extra.append((0, doc))

    pool = results[:DEEP_POOL] + extra
    if len(pool) < 2:
        return [doc for _, doc in pool[:limit]]

    # --- Re-puntuación por texto completo ----------------------------------
    # La pasada anterior solo mira título, materias y los primeros 400
    # caracteres. Eso basta cuando el asunto está en el título ("trabajo a
    # distancia"), pero no cuando vive en el articulado de un código: el
    # permiso por fallecimiento es el artículo 37 del Estatuto de los
    # Trabajadores, cuyo título no menciona permisos. Aquí se releen los
    # mejores candidatos y se puntúa su texto real.
    body: List[tuple] = []
    for _, doc in pool:
        puntos, _frags = _analyse_body(doc.get("_filepath", ""),
                                       tuple(terms), doc.get("title", ""))
        body.append((puntos, doc))

    body.sort(key=lambda x: -x[0])
    # Solo cuenta la evidencia textual FUERTE. Que un pasaje reúna dos
    # términos genéricos no dice nada: casi cualquier norma lo consigue, y
    # dejarlo entrar hundía resultados que la búsqueda por título ya acertaba.
    rank_texto = {id(doc): i for i, (puntos, doc) in enumerate(body)
                  if puntos >= BODY_MIN_SCORE}
    rank_clave = {id(doc): i for i, (_, doc) in enumerate(results[:DEEP_POOL])}

    # Fusión por rangos recíprocos de las tres señales, sin normalizar escalas
    # distintas. Título y articulado se complementan: el primero acierta con
    # las normas monográficas ("Ley de trabajo a distancia") y el segundo con
    # los códigos generales, donde el asunto vive en un artículo.
    def fusion(par: tuple) -> float:
        doc = par[1]
        clave = id(doc)
        s = 0.0
        if clave in rank_clave:
            s += 1.0 / (RRF_K + rank_clave[clave])
        if clave in rank_texto_completo:
            s += 1.0 / (RRF_K + rank_texto_completo[clave])
        if clave in rank_texto:
            s += BODY_WEIGHT / (RRF_K + rank_texto[clave])
        return -s

    return [doc for _, doc in sorted(pool, key=fusion)[:limit]]


# ---------------------------------------------------------------------------
# Búsqueda semántica (HuggingFace Embeddings)
# ---------------------------------------------------------------------------

def _cosine_sim(a: list, b: list) -> float:
    try:
        import numpy as np
        va, vb = np.array(a, dtype=float), np.array(b, dtype=float)
        d = np.linalg.norm(va) * np.linalg.norm(vb)
        return float(np.dot(va, vb) / d) if d > 1e-8 else 0.0
    except ImportError:
        dot  = sum(x * y for x, y in zip(a, b))
        na   = sum(x * x for x in a) ** 0.5
        nb   = sum(x * x for x in b) ** 0.5
        return dot / (na * nb) if na * nb > 1e-8 else 0.0


async def semantic_rerank(query: str, candidates: List[dict]) -> List[dict]:
    """Re-ordena candidatos por similitud semántica usando embeddings de HuggingFace."""
    if not candidates or not llm.HF_TOKEN:
        return candidates

    doc_texts = [
        f"{d.get('title', '')} | {' '.join((d.get('subjects') or [])[:6])}"
        for d in candidates
    ]
    all_texts = [query] + doc_texts

    loop = asyncio.get_event_loop()
    all_embs = await loop.run_in_executor(None, llm.embed_texts, all_texts)
    if not all_embs:
        return candidates

    query_emb = all_embs[0]
    scored = [(  _cosine_sim(query_emb, emb), doc)
              for emb, doc in zip(all_embs[1:], candidates)]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [doc for _, doc in scored]


# ---------------------------------------------------------------------------
# Consulta con el proveedor de IA activo
# ---------------------------------------------------------------------------

async def ask_claude(question: str, docs: List[dict],
                     precision: bool = False,
                     adjunto: Optional[dict] = None) -> llm.Result:
    if not llm.available():
        return llm.Result(text=None)

    n_docs   = CONTEXT_DOCS_PREC if precision else CONTEXT_DOCS
    n_frags  = FRAG_PER_DOC_PREC if precision else FRAG_PER_DOC
    width    = FRAG_WIDTH_PREC   if precision else FRAG_WIDTH
    selected = docs[:n_docs]

    # Releer los ficheros es E/S bloqueante: fuera del bucle de eventos.
    terms = _tokenise(_expand_query(question))
    loop = asyncio.get_event_loop()
    fragment_sets = await asyncio.gather(*[
        loop.run_in_executor(None, _relevant_fragments,
                             doc.get("_filepath", ""), terms, n_frags, width,
                             doc.get("title", ""))
        for doc in selected
    ])

    context_parts: List[str] = []
    for i, (doc, frags) in enumerate(zip(selected, fragment_sets), 1):
        # Si ningún pasaje coincide, al menos va el preámbulo de siempre.
        if frags:
            cuerpo = "\n\n".join(f"    {f}" for f in frags)
            etiqueta = "Pasajes relevantes del texto"
        else:
            cuerpo = f"    {doc.get('_preview', '')}"
            etiqueta = "Inicio del texto (sin coincidencias literales)"

        subjects = ", ".join(doc.get("subjects") or [])
        context_parts.append(
            f"[{i}] {doc.get('title', 'Sin título')}\n"
            f"    ID: {doc.get('identifier', '')}\n"
            f"    Tipo: {RANK_LABELS.get(doc.get('rank', ''), doc.get('rank', ''))}\n"
            f"    Fecha: {doc.get('publication_date', '')}\n"
            f"    Estado: {'Vigente' if doc.get('status') == 'in_force' else doc.get('status', '')}\n"
            f"    Departamento: {doc.get('department', '')}\n"
            f"    Materias: {subjects}\n"
            f"    Fuente: {doc.get('source', '')}\n\n"
            f"    {etiqueta}:\n{cuerpo}\n"
        )

    context = "\n---\n".join(context_parts)

    system_prompt = (
        "Eres un asistente jurídico especializado en legislación española. "
        "Tu misión es ayudar a ciudadanos, empresas y profesionales a entender "
        "la normativa vigente aplicable a sus situaciones de forma clara y práctica. "
        "Responde siempre en español. Cita los identificadores BOE cuando los menciones. "
        "Aplica la normativa legislativa aportada al caso planteado por el usuario de forma directa; "
        "no digas 'no hay información sobre la situación descrita' simplemente porque la ley "
        "no mencione el nombre o la entidad específica del usuario: explica cómo regula la "
        "ley los hechos expuestos. "
        "Prioriza Leyes Orgánicas, Leyes generales y Reales Decretos sobre convenios o resoluciones "
        "específicas de otros organismos. "
        "Sé conciso pero completo. Usa listas cuando sea útil.\n\n"
        "Cada norma se acompaña de pasajes literales de su texto, precedidos "
        "del encabezado del artículo al que pertenecen (por ejemplo "
        "'Artículo 37. Descanso semanal...'). Cita el artículo usando ese "
        "encabezado y reproduce el texto tal como aparece.\n"
        "Precisión ante todo: no cites un artículo cuyo encabezado no figure "
        "entre los pasajes aportados, ni entrecomilles texto que no esté ahí."
    )

    # Pasajes del documento aportado por el usuario, seleccionados por su
    # relación con la pregunta: una resolución de 20 páginas no cabe entera
    # en la ventana de 8.192 tokens del modelo local.
    bloque_adjunto = ""
    if adjunto and adjunto.get("texto"):
        cuerpo = _pasajes_adjunto(adjunto["texto"], terms)
        bloque_adjunto = (
            f"DOCUMENTO APORTADO POR LA PERSONA USUARIA "
            f"({adjunto.get('nombre', 'documento')}):\n{cuerpo}\n\n"
            "Este documento describe su caso concreto. Úsalo para los hechos, "
            "y la normativa de abajo para el fundamento jurídico.\n\n"
        )

    user_message = (
        f"Pregunta: {question}\n\n"
        f"{bloque_adjunto}"
        f"Normativa relevante encontrada, con pasajes literales de su texto:\n\n"
        f"{context}\n\n"
        "Por favor responde la pregunta basándote en estos documentos. "
        "Indica al final las referencias BOE utilizadas."
    )

    for d in selected:
        doc_id = d.get("identifier") or d.get("_filename")
        if doc_id:
            _record_reinforcement_feedback(doc_id, reward=1.0)

    return await asyncio.to_thread(
        llm.complete_ex, system_prompt, user_message, max_tokens=1500, precision=precision
    )


# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------

app = FastAPI(title="Consultor Legislativo España", version="1.0.0")

app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")


class QueryRequest(BaseModel):
    question: str
    region:   Optional[str] = None
    rank:     Optional[str] = None
    status:   Optional[str] = None
    semantic: bool = False
    # Antepone Anthropic para máxima precisión normativa. Si su cuota está
    # agotada se cae al modelo local automáticamente.
    precision: bool = False
    # Documento aportado por el usuario (ver /api/adjunto).
    adjunto_id: Optional[str] = None
    # Con adjunto se fuerza el modelo local salvo que se autorice la nube.
    permitir_nube: bool = False


class DocRequest(BaseModel):
    doc_type: str
    datos: dict
    use_ai: bool = True
    context_query: Optional[str] = None
    # Antepone Anthropic al redactar los fundamentos de derecho.
    precision: bool = False
    # Documento aportado (ver /api/adjunto): resolución que se recurre,
    # contrato, requerimiento… De él salen los hechos del escrito.
    adjunto_id: Optional[str] = None
    permitir_nube: bool = False


@app.get("/", response_class=FileResponse)
def root():
    return FileResponse(str(BASE_DIR / "static" / "index.html"))


@app.get("/api/status")
def api_status():
    ready = _index_ready.is_set()
    prov  = llm.provider()
    return {
        "ready":        ready,
        "total_docs":   len(_index) if ready else _index_loaded,
        "total_files":  _index_total,
        "corpus_path":  str(CORPUS_PATH),
        "index_error":  _index_error or "",
        "has_ai":       llm.available(),
        "ai_provider":  prov,
        "ai_model":     llm.active_model(),
        # Modo precisión: qué proveedor encabezaría la cadena y si está
        # disponible, para que la interfaz pueda ofrecer la opción o no.
        "has_precision":       bool(llm.ANTHROPIC_KEY),
        "precision_provider":  llm.provider(precision=True),
        "precision_model":     llm.active_model(precision=True),
        "ai_chain":            llm.chain(),
        "precision_chain":     llm.chain(precision=True),
        "has_semantic": bool(llm.HF_TOKEN),
        "embed_model":  llm.HF_EMBED_MODEL if llm.HF_TOKEN else "",
        "regions":      REGIONS,
        "ranks":        RANK_LABELS,
    }


@app.get("/api/search")
async def api_search(
    q:        str           = Query(..., min_length=2),
    region:   Optional[str] = None,
    rank:     Optional[str] = None,
    status:   Optional[str] = None,
    limit:    int           = Query(default=20, le=50),
    semantic: bool          = False,
    deep:     bool          = True,
):
    if not _index_ready.is_set():
        raise HTTPException(503, "Índice en construcción, espera unos segundos.")

    pool = SEMANTIC_POOL if (semantic and llm.HF_TOKEN) else limit
    loop = asyncio.get_event_loop()
    docs = await loop.run_in_executor(
        None, lambda: search(q, region=region, rank_filter=rank, status_filter=status, limit=pool, deep=deep)
    )

    if semantic and llm.HF_TOKEN and docs:
        docs = await semantic_rerank(q, docs)
        docs = docs[:limit]

    return {
        "query":    q,
        "total":    len(docs),
        "semantic": semantic and bool(llm.HF_TOKEN),
        "results":  [_serialise(d) for d in docs],
    }


@app.post("/api/adjunto")
async def api_adjunto(file: UploadFile = File(...)):
    """Sube un documento, extrae su texto y devuelve un identificador.

    Ese identificador se pasa luego a /api/ask o /api/generar. El texto vive
    solo en memoria y caduca en una hora: son documentos con datos personales
    y no hay motivo para dejarlos escritos en disco.
    """
    datos = await file.read()
    if not datos:
        raise HTTPException(400, "El fichero está vacío.")
    if len(datos) > ADJUNTO_MAX_BYTES:
        raise HTTPException(413, f"El fichero supera el límite de "
                                 f"{ADJUNTO_MAX_BYTES // (1024*1024)} MB.")

    try:
        texto = _extraer_texto(file.filename or "", datos)
    except ValueError as exc:
        raise HTTPException(415, str(exc))
    except Exception as exc:
        raise HTTPException(422, f"No se pudo leer el documento: {exc}")

    texto = texto.strip()
    if not texto:
        raise HTTPException(422, "No se ha podido extraer texto del documento. "
                                 "Comprueba que el archivo o imagen escaneada no esté protegido "
                                 "y tenga suficiente nitidez y resolución para el motor OCR.")

    recortado = len(texto) > ADJUNTO_MAX_CHARS
    texto = texto[:ADJUNTO_MAX_CHARS]

    _purgar_adjuntos()
    adjunto_id = uuid.uuid4().hex
    with _adjuntos_lock:
        _adjuntos[adjunto_id] = {
            "texto": texto,
            "nombre": file.filename or "documento",
            "ts": time.time(),
        }

    terminos = _terminos_salientes(texto)
    return {
        "adjunto_id":  adjunto_id,
        "nombre":      file.filename,
        "caracteres":  len(texto),
        "recortado":   recortado,
        "caduca_en":   ADJUNTO_TTL,
        "terminos":    terminos,
        "vista_previa": texto[:400],
        # Aviso explícito: con IA en la nube el contenido sale del equipo.
        "aviso_privacidad": (
            "El texto se conserva solo en memoria y caduca en una hora. "
            "Las consultas con adjunto se atienden con el modelo local para "
            "que el contenido no salga de este equipo; usa "
            "\"permitir_nube\": true si aceptas enviarlo a HuggingFace o "
            "Anthropic."
        ),
    }


@app.post("/api/ask")
async def api_ask(req: QueryRequest):
    if not _index_ready.is_set():
        raise HTTPException(503, "Índice en construcción, espera unos segundos.")

    # --- Documento aportado -------------------------------------------------
    adjunto = None
    aviso = None
    if req.adjunto_id:
        _purgar_adjuntos()
        with _adjuntos_lock:
            adjunto = _adjuntos.get(req.adjunto_id)

        if adjunto is None:
            raise HTTPException(404, "El adjunto no existe o ha caducado. "
                                     "Vuelve a subirlo.")

    # La búsqueda se enriquece con los términos que caracterizan al adjunto:
    # una resolución sancionadora de tráfico lleva a la normativa de tráfico
    # aunque la pregunta sea tan escueta como "¿puedo recurrir esto?".
    consulta = req.question
    if adjunto:
        extra = _terminos_salientes(adjunto["texto"], tope=10)
        if extra:
            consulta = f"{req.question} {' '.join(extra)}"

    # Con adjunto, su contenido no sale del equipo salvo autorización expresa:
    # puede llevar nombre, DNI o datos bancarios, propios o de terceros.
    precision = req.precision
    if adjunto and not req.permitir_nube:
        if llm.provider() == "local":
            precision = False            # el local encabeza la cadena normal
        else:
            aviso = ("No hay modelo local disponible, así que el adjunto no se "
                     "ha usado para no enviarlo a la nube. Vuelve a preguntar "
                     "con \"permitir_nube\": true si autorizas el envío.")
            adjunto = None

    use_semantic = req.semantic and bool(llm.HF_TOKEN)
    pool = SEMANTIC_POOL if use_semantic else 10
    loop = asyncio.get_event_loop()
    docs = await loop.run_in_executor(
        None, lambda: search(consulta, region=req.region, rank_filter=req.rank, status_filter=req.status, limit=pool)
    )

    if use_semantic and docs:
        docs = await semantic_rerank(consulta, docs)

    if docs:
        rank_weight = {
            "ley_organica": 5, "ley": 4, "real_decreto_ley": 4,
            "real_decreto": 3, "decreto": 3, "orden": 2, "resolucion": 1
        }
        docs.sort(key=lambda d: (rank_weight.get(d.get("rank", ""), 0), d.get("publication_date") or ""), reverse=True)

    result = await ask_claude(req.question, docs[:8], precision=precision,
                              adjunto=adjunto)

    return {
        "question":    req.question,
        "ai_answer":   result.text,
        # Se informa del proveedor que respondió de verdad, no del que tocaba
        # por orden: si Anthropic agotó cuota, aquí saldrá 'local'.
        "ai_provider": result.provider,
        "ai_model":    result.model,
        "precision":   precision,
        "ai_fallback": [{"provider": p, "reason": r} for p, r in result.attempts],
        "semantic":    bool(llm.HF_TOKEN),
        "adjunto":     adjunto["nombre"] if adjunto else None,
        "aviso":       aviso,
        "sources":     [_serialise(d) for d in docs[:8]],
    }


@app.get("/api/recent")
def api_recent(region: Optional[str] = None, limit: int = 15):
    if not _index_ready.is_set():
        raise HTTPException(status_code=503, detail="Índice en construcción.")

    docs = [
        d for d in _index
        if (not region or d.get("_region") == region)
        and d.get("publication_date")
    ]
    docs.sort(key=lambda d: d.get("publication_date") or "", reverse=True)
    return {"results": [_serialise(d) for d in docs[:limit]]}


@app.get("/api/doc/{identifier}")
def api_doc(identifier: str):
    if not _index_ready.is_set():
        raise HTTPException(status_code=503, detail="Índice en construcción.")

    for doc in _index:
        if doc.get("identifier") == identifier or doc.get("_filename") == f"{identifier}.md":
            fp = Path(doc["_filepath"])
            try:
                full_text = fp.read_text(errors="ignore")
                end = full_text.find("---", 3)
                content = full_text[end + 3:].strip() if end != -1 else full_text
                return {**_serialise(doc), "full_content": content[:8000]}
            except OSError:
                return _serialise(doc)

    raise HTTPException(status_code=404, detail="Documento no encontrado.")


# ---------------------------------------------------------------------------
# Generación de documentos
# ---------------------------------------------------------------------------

@app.get("/api/doc-types")
def api_doc_types():
    """Catálogo de tipos de documento generables y sus campos."""
    return {
        "has_pdf": documents.LATEX_ENGINE is not None,
        "has_ai": llm.available(),
        "types": documents.doc_types_catalog(),
    }


@app.post("/api/generar")
async def api_generar(req: DocRequest):
    if req.doc_type not in documents.DOC_TYPES:
        raise HTTPException(status_code=400, detail="Tipo de documento no válido.")

    cfg = documents.DOC_TYPES[req.doc_type]
    faltan = [
        f["label"] for f in cfg["fields"]
        if f.get("required") and not str(req.datos.get(f["name"], "")).strip()
    ]
    if faltan:
        raise HTTPException(
            status_code=422,
            detail="Faltan campos obligatorios: " + ", ".join(faltan),
        )

    # --- Documento aportado -------------------------------------------------
    adjunto = None
    aviso = None
    if req.adjunto_id:
        with _adjuntos_lock:
            adjunto = _adjuntos.get(req.adjunto_id)
        if adjunto is None:
            raise HTTPException(404, "El adjunto no existe o ha caducado. "
                                     "Vuelve a subirlo.")

    # Igual que en /api/ask: el contenido no sale del equipo sin permiso.
    precision = req.precision
    if adjunto and not req.permitir_nube:
        if llm.provider() == "local":
            precision = False
        else:
            aviso = ("No hay modelo local disponible, así que el adjunto no se "
                     "ha usado para no enviarlo a la nube. Repite con "
                     "\"permitir_nube\": true si autorizas el envío.")
            adjunto = None

    # Buscamos la normativa relacionada con el asunto/hechos para (a) dar contexto
    # a la IA y (b) citarla SIEMPRE en el documento como "Normativa aplicable",
    # aunque la IA deje vacíos los fundamentos o esté desactivada.
    query = req.context_query or req.datos.get("asunto") or req.datos.get("hechos", "")
    # El adjunto es la mejor fuente para localizar la normativa: la resolución
    # que se recurre nombra la materia mucho mejor que un asunto de dos palabras.
    if adjunto:
        extra = _terminos_salientes(adjunto["texto"], tope=10)
        if extra:
            query = f"{query} {' '.join(extra)}".strip()

    normativa = search(query, status_filter="in_force", limit=16) if _index_ready.is_set() and query else []
    if llm.HF_TOKEN and normativa:
        normativa = await semantic_rerank(query, normativa)
    normativa = normativa[:8]

    ai_sections = None
    if req.use_ai and llm.available():
        terms = _tokenise(_expand_query(query))

        pasajes = _pasajes_adjunto(adjunto["texto"], terms) if adjunto else ""

        # Pasajes literales del articulado de cada norma, igual que en
        # /api/ask. Sin ellos el modelo solo ve títulos y deja los
        # fundamentos en blanco por no poder citar nada concreto.
        bloques: List[str] = []
        loop = asyncio.get_event_loop()
        frags = await asyncio.gather(*[
            loop.run_in_executor(None, _relevant_fragments,
                                 d.get("_filepath", ""), terms,
                                 FRAG_PER_DOC, FRAG_WIDTH, d.get("title", ""))
            for d in normativa[:CONTEXT_DOCS]
        ]) if normativa else []
        for i, (d, fs) in enumerate(zip(normativa[:CONTEXT_DOCS], frags), 1):
            cabecera = (f"[{i}] {d.get('title', 'Sin título')} "
                        f"(BOE {d.get('identifier', '')}, "
                        f"{d.get('publication_date', '')})")
            cuerpo = "\n".join(f"    {f}" for f in fs) if fs else ""
            bloques.append(f"{cabecera}\n{cuerpo}" if cuerpo else cabecera)

        ai_sections = await documents.draft_with_claude(
            req.doc_type, req.datos, normativa, precision=precision,
            adjunto_texto=pasajes,
            adjunto_nombre=adjunto["nombre"] if adjunto else "",
            contexto_normativa="\n\n".join(bloques))

    result = await asyncio.to_thread(
        documents.generate, req.doc_type, req.datos, ai_sections, normativa
    )
    if isinstance(result, dict):
        result["adjunto"] = adjunto["nombre"] if adjunto else None
        result["aviso"] = aviso
    return result


@app.get("/api/generar/{doc_id}.{fmt}")
def api_download(doc_id: str, fmt: str):
    if fmt not in ("tex", "pdf"):
        raise HTTPException(status_code=400, detail="Formato no válido.")
    path = documents.file_path(doc_id, fmt)
    if not path:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    media = "application/pdf" if fmt == "pdf" else "application/x-tex"
    return FileResponse(
        str(path),
        media_type=media,
        filename=f"documento.{fmt}",
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _serialise(doc: dict) -> dict:
    """Return a JSON-safe subset of a document record."""
    return {
        "identifier": doc.get("identifier", ""),
        "title": doc.get("title", "Sin título"),
        "rank": doc.get("rank", ""),
        "rank_label": RANK_LABELS.get(doc.get("rank", ""), doc.get("rank", "")),
        "publication_date": doc.get("publication_date", ""),
        "last_updated": doc.get("last_updated", ""),
        "status": doc.get("status", ""),
        "department": doc.get("department", ""),
        "subjects": doc.get("subjects") or [],
        "source": doc.get("source", ""),
        "pdf_url": doc.get("pdf_url") or doc.get("url_pdf", ""),
        "scope": doc.get("scope", ""),
        "region": REGIONS.get(doc.get("_region", ""), doc.get("_region", "")),
        "region_code": doc.get("_region", ""),
        "preview": doc.get("_preview", ""),
        "official_number": doc.get("official_number", ""),
        "alerts": doc.get("alerts", ""),
    }


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=False)
